"""
Convert Chapter 4 Markdown to Word Document
Properly formats tables, headings, and content for thesis formatting
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def add_heading(doc, text, level):
    """Add formatted heading"""
    heading = doc.add_heading(text, level=level)
    if level == 1:
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.runs[0]
        run.font.size = Pt(16)
        run.bold = True
    elif level == 2:
        run = heading.runs[0]
        run.font.size = Pt(14)
        run.bold = True
    elif level == 3:
        run = heading.runs[0]
        run.font.size = Pt(12)
        run.bold = True
    return heading

def parse_table(lines):
    """Parse markdown table into structured data"""
    rows = []
    for line in lines:
        if '|' in line and not line.strip().startswith('|---'):
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if cells and any(cell for cell in cells):  # Skip empty rows
                rows.append(cells)
    return rows

def add_table_to_doc(doc, table_data):
    """Add formatted table to document"""
    if not table_data or len(table_data) < 2:
        return
    
    # Create table
    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
    table.style = 'Light Grid Accent 1'
    
    # Populate table
    for i, row_data in enumerate(table_data):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_data
            
            # Format header row
            if i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(11)
            else:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
    
    return table

def clean_markdown_formatting(text):
    """Remove markdown formatting like ** for bold"""
    # Remove bold markers
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    # Remove italic markers
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    # Remove code markers
    text = re.sub(r'`([^`]+)`', r'\1', text)
    return text

def convert_markdown_to_word(md_file, output_file):
    """Convert markdown file to Word document"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    in_table = False
    table_lines = []
    in_code_block = False
    code_lines = []
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Handle code blocks
        if line.startswith('```'):
            if in_code_block:
                # End code block
                if code_lines:
                    code_text = '\n'.join(code_lines)
                    p = doc.add_paragraph(code_text)
                    p.style = 'No Spacing'
                    for run in p.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                code_lines = []
                in_code_block = False
            else:
                # Start code block
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # Handle tables
        if '|' in line and line.strip():
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        elif in_table:
            # End of table
            if table_lines:
                table_data = parse_table(table_lines)
                add_table_to_doc(doc, table_data)
                doc.add_paragraph()  # Add space after table
            in_table = False
            table_lines = []
        
        # Handle headings
        if line.startswith('# '):
            heading_text = line[2:].strip()
            add_heading(doc, heading_text, level=1)
        elif line.startswith('## '):
            heading_text = line[3:].strip()
            add_heading(doc, heading_text, level=2)
        elif line.startswith('### '):
            heading_text = line[4:].strip()
            add_heading(doc, heading_text, level=3)
        elif line.startswith('#### '):
            heading_text = line[5:].strip()
            add_heading(doc, heading_text, level=4)
        # Handle bullet points
        elif line.strip().startswith('- '):
            text = clean_markdown_formatting(line.strip()[2:])
            p = doc.add_paragraph(text, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.5)
        # Handle numbered lists
        elif re.match(r'^\d+\.\s', line.strip()):
            text = clean_markdown_formatting(re.sub(r'^\d+\.\s', '', line.strip()))
            p = doc.add_paragraph(text, style='List Number')
            p.paragraph_format.left_indent = Inches(0.5)
        # Handle bold text in lines starting with **
        elif line.strip().startswith('**') and line.strip().endswith('**'):
            text = clean_markdown_formatting(line.strip())
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
        # Handle horizontal rules
        elif line.strip() == '---':
            doc.add_paragraph('_' * 80)
        # Handle regular paragraphs
        elif line.strip():
            text = clean_markdown_formatting(line.strip())
            p = doc.add_paragraph(text)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # Handle empty lines
        else:
            if i > 0 and lines[i-1].strip():  # Only add space if previous line had content
                doc.add_paragraph()
        
        i += 1
    
    # Save document
    doc.save(output_file)
    print(f"✅ Document saved to: {output_file}")

if __name__ == "__main__":
    md_file = "CHAPTER_4_RESULTS_AND_DISCUSSION.md"
    output_file = "CHAPTER_4_RESULTS_AND_DISCUSSION.docx"
    
    print(f"Converting {md_file} to {output_file}...")
    convert_markdown_to_word(md_file, output_file)
    print("Conversion complete!")
